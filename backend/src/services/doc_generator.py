from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from src.models.schemas import ResumeData


def _add_bottom_border(paragraph):
    """Add a thin bottom border (horizontal rule) to a paragraph."""
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_run_font(run, name="Times New Roman", size=10.5, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _add_section_heading(doc, title: str):
    """ALL-CAPS bold heading with horizontal rule underneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title.upper())
    _set_run_font(run, bold=True)
    _add_bottom_border(p)
    return p


def _add_two_column_line(doc, left: str, right: str, bold=False, italic=False):
    """Add a line with left text and right-aligned text using a tab stop."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Right-aligned tab stop at page width
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run_left = p.add_run(left)
    _set_run_font(run_left, bold=bold, italic=italic)
    p.add_run("\t")
    run_right = p.add_run(right)
    _set_run_font(run_right, bold=bold, italic=italic)
    return p


def _add_bullet(doc, text: str):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.clear()
    run = p.add_run(text)
    _set_run_font(run)
    return p


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function that places a hyperlink within a paragraph object.
    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color if it is given
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)
    else:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def generate_docx_bytes(resume: ResumeData) -> bytes:
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ── Name ─────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(resume.full_name)
    _set_run_font(run, size=18, bold=True)

    # ── Contact line ─────────────────────────────────────────────
    contact_items = []
    if resume.phone:
        contact_items.append((resume.phone, None))
    if resume.email:
        contact_items.append((resume.email, None))
    if resume.linkedin:
        contact_items.append(("LinkedIn", resume.linkedin))
    if resume.location:
        contact_items.append((resume.location, None))
        
    if contact_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        
        for i, (text, url) in enumerate(contact_items):
            if i > 0:
                run = p.add_run(" | ")
                _set_run_font(run)
                
            if url:
                add_hyperlink(p, url, text)
            else:
                run = p.add_run(text)
                _set_run_font(run)

    # ── Summary ──────────────────────────────────────────────────
    if resume.summary:
        _add_section_heading(doc, "Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(resume.summary)
        _set_run_font(run)

    # ── Experience ───────────────────────────────────────────────
    if resume.experience:
        _add_section_heading(doc, "Experience")
        for exp in resume.experience:
            _add_two_column_line(doc, exp.company or "", exp.location or "", bold=True)
            dates = f"{exp.start_date or ''} - {exp.end_date or 'Present'}"
            _add_two_column_line(doc, exp.job_title or "", dates, italic=True)
            for r in (exp.responsibilities or []):
                _add_bullet(doc, r)

    # ── Projects ─────────────────────────────────────────────────
    if resume.projects:
        _add_section_heading(doc, "Projects")
        for proj in resume.projects:
            techs = ", ".join(proj.technologies or [])
            link = proj.link or ""
            
            # 1. Title Line
            if link and (link.startswith("http://") or link.startswith("https://")):
                # Left: Title, Right: "GitHub Link" (Hyperlink)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                
                # Title (Left)
                run_title = p.add_run(proj.title or "")
                _set_run_font(run_title, bold=True)
                
                # Tab
                p.add_run("\t")
                
                # Link (Right) - "GitHub Link"
                # Add hyperlink part
                part = doc.part
                r_id = part.relate_to(link, "hyperlink", is_external=True)
                
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                
                new_run = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                
                # Style for link (Blue + Underline + Bold to match title line style)
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "0000FF")
                rPr.append(color)
                
                u = OxmlElement("w:u")
                u.set(qn("w:val"), "single")
                rPr.append(u)
                
                b = OxmlElement("w:b")
                rPr.append(b)
                
                # Font size 9.5pt (19 half-points) - decreased by 1 from base 10.5
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "19")
                rPr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), "19")
                rPr.append(szCs)
                
                new_run.append(rPr)
                new_run.text = "GitHub Link"
                
                hyperlink.append(new_run)
                p._p.append(hyperlink)

            else:
                 # Just title or Title vs Link text
                 _add_two_column_line(doc, proj.title or "", link, bold=True)
            
            # 2. Technologies Line
            if techs:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(f"Technologies: {techs}")
                _set_run_font(run, italic=True)

            # 3. Description
            if proj.description:
                _add_bullet(doc, proj.description)

    # ── Skills ───────────────────────────────────────────────────
    if resume.skills:
        _add_section_heading(doc, "Skills")
        for skill in resume.skills:
             # Clean leading bullets if present (to avoid double bulleting)
            text = skill.lstrip(" •●-")
            if not text: continue
            _add_bullet(doc, text)

    # ── Education ────────────────────────────────────────────────
    if resume.education:
        _add_section_heading(doc, "Education")
        for ed in resume.education:
            _add_two_column_line(doc, ed.institution or "", ed.location or "", bold=True)
            _add_two_column_line(doc, ed.degree or "", ed.graduation_date or "", italic=True)

    # ── Certifications ───────────────────────────────────────────
    if resume.certifications:
        _add_section_heading(doc, "Certifications")
        for cert in resume.certifications:
            _add_bullet(doc, cert)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

